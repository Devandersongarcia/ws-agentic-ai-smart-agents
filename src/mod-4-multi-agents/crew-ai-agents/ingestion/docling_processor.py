"""
Docling-based document processor for extracting structured content.
Handles PDFs, JSON, CSV, and DOCX files with structure preservation.
"""

import json
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

import pandas as pd
from docling.document_converter import DocumentConverter, PdfPipelineOptions
from docling.datamodel.base_models import InputFormat
from docling.datamodel.document import ConversionResult
from docling.datamodel.pipeline_options import TableFormerMode
from docx import Document as DocxDocument

from config import get_settings
from database.models import Document, DocumentMetadata, DocumentType, CuisineType


class DoclingProcessor:
    """
    Processes documents using IBM's Docling library.
    Preserves document structure including tables, formulas, and layouts.
    """

    def __init__(self):
        """Initialize Docling processor with configuration."""
        self.settings = get_settings()
        self._setup_pipeline_options()
        self.converter = DocumentConverter(pipeline_options=self.pipeline_options)

    def _setup_pipeline_options(self) -> None:
        """Configure Docling pipeline options for optimal extraction."""
        self.pipeline_options = PdfPipelineOptions()
        self.pipeline_options.do_ocr = self.settings.docling_ocr_enabled
        self.pipeline_options.do_table_structure = True
        self.pipeline_options.table_structure_options.mode = TableFormerMode.ACCURATE
        
        if self.settings.docling_use_gpu:
            self.pipeline_options.accelerator = "cuda"
        else:
            self.pipeline_options.accelerator = "cpu"

    def process_file(self, file_path: Path) -> Document:
        """
        Process a single file and extract structured content.
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            Processed Document object
            
        Raises:
            ValueError: If file type is not supported
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension == ".pdf":
            return self._process_pdf(file_path)
        elif file_extension == ".json":
            return self._process_json(file_path)
        elif file_extension == ".csv":
            return self._process_csv(file_path)
        elif file_extension == ".docx":
            return self._process_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

    def _process_pdf(self, file_path: Path) -> Document:
        """
        Process PDF file using Docling with table preservation.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Processed Document with extracted content
        """
        result: ConversionResult = self.converter.convert(str(file_path))
        
        markdown_content = result.document.export_to_markdown()
        
        cuisine_type = self._extract_cuisine_type(file_path.stem)
        
        metadata = DocumentMetadata(
            source=str(file_path),
            doc_type=DocumentType.PDF,
            cuisine_type=cuisine_type,
            total_pages=len(result.document.pages) if hasattr(result.document, 'pages') else 1,
            extra={
                "has_tables": self._check_for_tables(result),
                "has_images": self._check_for_images(result),
            }
        )
        
        doc = Document(
            title=file_path.stem,
            content=markdown_content,
            metadata=metadata
        )
        
        return doc

    def _process_json(self, file_path: Path) -> Document:
        """
        Process JSON file containing restaurant data.
        
        Args:
            file_path: Path to JSON file
            
        Returns:
            Processed Document with structured content
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        content = self._json_to_markdown(data)
        
        metadata = DocumentMetadata(
            source=str(file_path),
            doc_type=DocumentType.JSON,
            extra={"record_count": len(data) if isinstance(data, list) else 1}
        )
        
        doc = Document(
            title=file_path.stem,
            content=content,
            metadata=metadata
        )
        
        return doc

    def _process_csv(self, file_path: Path) -> Document:
        """
        Process CSV file containing coupon data.
        
        Args:
            file_path: Path to CSV file
            
        Returns:
            Processed Document with tabular content
        """
        df = pd.read_csv(file_path)
        
        content = self._dataframe_to_markdown(df)
        
        metadata = DocumentMetadata(
            source=str(file_path),
            doc_type=DocumentType.CSV,
            extra={
                "row_count": len(df),
                "column_count": len(df.columns),
                "columns": df.columns.tolist()
            }
        )
        
        doc = Document(
            title=file_path.stem,
            content=content,
            metadata=metadata
        )
        
        return doc

    def _process_docx(self, file_path: Path) -> Document:
        """
        Process DOCX file containing allergen guidelines.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Processed Document with formatted content
        """
        docx = DocxDocument(file_path)
        
        content_parts = []
        
        for element in docx.element.body:
            if element.tag.endswith('p'):
                para = element.text
                if para.strip():
                    content_parts.append(para.strip())
            elif element.tag.endswith('tbl'):
                table_content = self._extract_docx_table(element)
                if table_content:
                    content_parts.append(table_content)
        
        content = "\n\n".join(content_parts)
        
        metadata = DocumentMetadata(
            source=str(file_path),
            doc_type=DocumentType.DOCX,
            extra={
                "paragraph_count": len(docx.paragraphs),
                "table_count": len(docx.tables)
            }
        )
        
        doc = Document(
            title=file_path.stem,
            content=content,
            metadata=metadata
        )
        
        return doc

    def _extract_cuisine_type(self, filename: str) -> Optional[CuisineType]:
        """
        Extract cuisine type from filename.
        
        Args:
            filename: Name of the file
            
        Returns:
            CuisineType if found, None otherwise
        """
        filename_lower = filename.lower()
        for cuisine in CuisineType:
            if cuisine.value in filename_lower:
                return cuisine
        return None

    def _check_for_tables(self, result: ConversionResult) -> bool:
        """
        Check if document contains tables.
        
        Args:
            result: Docling conversion result
            
        Returns:
            True if tables are present
        """
        if hasattr(result.document, 'tables'):
            return len(result.document.tables) > 0
        return False

    def _check_for_images(self, result: ConversionResult) -> bool:
        """
        Check if document contains images.
        
        Args:
            result: Docling conversion result
            
        Returns:
            True if images are present
        """
        if hasattr(result.document, 'images'):
            return len(result.document.images) > 0
        return False

    def _json_to_markdown(self, data: Union[Dict, List]) -> str:
        """
        Convert JSON data to markdown format.
        
        Args:
            data: JSON data to convert
            
        Returns:
            Markdown formatted string
        """
        if isinstance(data, list):
            parts = []
            for i, item in enumerate(data, 1):
                parts.append(f"## Record {i}\n")
                parts.append(self._dict_to_markdown(item))
            return "\n\n".join(parts)
        else:
            return self._dict_to_markdown(data)

    def _dict_to_markdown(self, data: Dict[str, Any], level: int = 0) -> str:
        """
        Convert dictionary to markdown format.
        
        Args:
            data: Dictionary to convert
            level: Indentation level
            
        Returns:
            Markdown formatted string
        """
        lines = []
        indent = "  " * level
        
        for key, value in data.items():
            if isinstance(value, dict):
                lines.append(f"{indent}**{key}:**")
                lines.append(self._dict_to_markdown(value, level + 1))
            elif isinstance(value, list):
                lines.append(f"{indent}**{key}:**")
                for item in value:
                    if isinstance(item, dict):
                        lines.append(self._dict_to_markdown(item, level + 1))
                    else:
                        lines.append(f"{indent}  - {item}")
            else:
                lines.append(f"{indent}**{key}:** {value}")
        
        return "\n".join(lines)

    def _dataframe_to_markdown(self, df: pd.DataFrame) -> str:
        """
        Convert DataFrame to markdown table.
        
        Args:
            df: DataFrame to convert
            
        Returns:
            Markdown table string
        """
        headers = "| " + " | ".join(df.columns) + " |"
        separator = "| " + " | ".join(["---"] * len(df.columns)) + " |"
        
        rows = []
        for _, row in df.iterrows():
            row_str = "| " + " | ".join(str(v) for v in row.values) + " |"
            rows.append(row_str)
        
        table_parts = [headers, separator] + rows
        return "\n".join(table_parts)

    def _extract_docx_table(self, table_element) -> str:
        """
        Extract table content from DOCX element.
        
        Args:
            table_element: Table XML element
            
        Returns:
            Markdown formatted table string
        """
        rows = []
        for row in table_element.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr'):
            cells = []
            for cell in row.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc'):
                cell_text = " ".join(p.text for p in cell.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p') if p.text)
                cells.append(cell_text.strip())
            if cells:
                rows.append(cells)
        
        if not rows:
            return ""
        
        max_cols = max(len(row) for row in rows)
        normalized_rows = []
        for row in rows:
            while len(row) < max_cols:
                row.append("")
            normalized_rows.append(row)
        
        header = "| " + " | ".join(normalized_rows[0]) + " |" if normalized_rows else ""
        separator = "| " + " | ".join(["---"] * max_cols) + " |"
        body_rows = []
        for row in normalized_rows[1:]:
            body_rows.append("| " + " | ".join(row) + " |")
        
        return "\n".join([header, separator] + body_rows) if header else ""

    def process_directory(self, directory: Path, file_pattern: str = "*") -> List[Document]:
        """
        Process all matching files in a directory.
        
        Args:
            directory: Directory path to process
            file_pattern: Glob pattern for file matching
            
        Returns:
            List of processed Documents
        """
        documents = []
        for file_path in directory.glob(file_pattern):
            if file_path.is_file():
                try:
                    doc = self.process_file(file_path)
                    documents.append(doc)
                    print(f"✓ Processed: {file_path.name}")
                except Exception as e:
                    print(f"✗ Failed to process {file_path.name}: {e}")
        
        return documents